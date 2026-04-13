import streamlit as st
import numpy as np
import io
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def create_docx(data):
    doc = Document()
    title = doc.add_heading('화성 탐사선 궤도 설계 최종 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run(f"소속: 사곡고등학교 | 성명: {data['name']} | 일자: {datetime.now().strftime('%Y-%m-%d')}").italic = True
    doc.add_heading('1. 궤도 설계 제원', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    params = [('장반경 (a)', f"{data['a_km']} km"), ('이심률 (e)', f"{data['e']}"), ('근일점 고도', f"{data['rp_alt']:.1f} km"), ('공전 주기 (T)', f"{data['T_h']:.2f} 시간")]
    for i, (label, val) in enumerate(params):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = val
    doc.add_heading('2. 탐구 분석 및 결론', level=1)
    questions = [("질문 1: 고해상도 촬영 궤도의 트레이드오프 분석", data.get('ans1', '')), ("질문 2: 케플러 제2법칙 적용 분석", data.get('ans2', '')), ("질문 3: 미션 주기 및 운용 전략", data.get('ans3', ''))]
    for q, ans in questions:
        doc.add_paragraph(q, style='List Bullet').bold = True
        doc.add_paragraph(f"답변: {ans if ans else '________________'}")
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

def run_report_ui():
    st.title("📝 궤도 설계 보고서 작성")
    if not DOCX_AVAILABLE:
        st.error("python-docx 라이브러리가 필요합니다.")
        return
    with st.form("report_form"):
        name = st.text_input("연구원 성명")
        a_km = st.number_input("설계 장반경 (km)", value=10000)
        e = st.number_input("설계 이심률", value=0.4)
        ans1 = st.text_area("탐구 1")
        ans2 = st.text_area("탐구 2")
        ans3 = st.text_area("탐구 3")
        if st.form_submit_button("보고서 생성"):
            rp_alt = a_km * (1 - e) - 3389.5
            T_h = (2 * np.pi * np.sqrt((a_km*1000)**3 / (6.674e-11 * 6.39e23))) / 3600
            docx_file = create_docx({'name': name, 'a_km': a_km, 'e': e, 'rp_alt': rp_alt, 'T_h': T_h, 'ans1': ans1, 'ans2': ans2, 'ans3': ans3})
            st.download_button("📂 DOCX 다운로드", data=docx_file, file_name=f"KASA_Report_{name}.docx")

if __name__ == "__main__":
    run_report_ui()
