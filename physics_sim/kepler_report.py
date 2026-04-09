import streamlit as st
import numpy as np
import io
from datetime import datetime

# ── python-docx 임포트 (없으면 안내) ───────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ── 물리 상수 ──────────────────────────────────────────────────────
G = 6.674e-11
PLANETS = {
    "화성 (Mars)":    {"M": 6.39e23,  "R_km": 3389.5},
    "지구 (Earth)":   {"M": 5.972e24, "R_km": 6371.0},
    "달 (Moon)":      {"M": 7.342e22, "R_km": 1737.4},
    "목성 (Jupiter)": {"M": 1.898e27, "R_km": 71492.0},
}

def calc(a_km, e, M):
    a = a_km * 1000
    T  = 2 * np.pi * np.sqrt(a**3 / (G * M))
    rp = a * (1 - e)
    ra = a * (1 + e)
    vp = np.sqrt(G * M * (2/rp - 1/a))
    va = np.sqrt(G * M * (2/ra - 1/a))
    return {"T_h": T/3600, "T_d": T/86400,
            "rp_km": rp/1000, "ra_km": ra/1000,
            "vp_kms": vp/1000, "va_kms": va/1000}

# ── DOCX 생성 함수 ─────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=(30, 41, 59)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_section_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def create_docx(form_data: dict, orbit: dict) -> bytes:
    doc = Document()

    # ── 페이지 여백 설정 ──
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── 기본 스타일 ──
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10.5)

    # ── 제목 헤더 ──────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("🚀 화성 탐사선 '다누리 2호' 궤도 설계 연구보고서")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138)
    run.font.name = '맑은 고딕'

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("케플러 법칙을 적용한 화성 탐사 궤도 설계 및 물리학적 검증")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    sub_run.font.name = '맑은 고딕'

    doc.add_paragraph()

    # ── 학생 정보 테이블 ───────────────────────────────────────
    info_table = doc.add_table(rows=2, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'

    headers = ["학번", "이름", "제출일", "선택 천체"]
    values  = [
        form_data.get("student_id", ""),
        form_data.get("student_name", ""),
        form_data.get("submit_date", ""),
        form_data.get("planet", "화성 (Mars)")
    ]
    for i, (h, v) in enumerate(zip(headers, values)):
        hcell = info_table.cell(0, i)
        vcell = info_table.cell(1, i)
        hcell.text = h
        vcell.text = v
        set_cell_bg(hcell, "1E3A8A")
        for run in hcell.paragraphs[0].runs:
            run.font.bold  = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = '맑은 고딕'
        vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_section_divider(doc)
    doc.add_paragraph()

    # ── 1. 궤도 설계 파라미터 ──────────────────────────────────
    add_heading(doc, "1. 설계 궤도 파라미터", level=1, color=(30, 58, 138))
    doc.add_paragraph()

    param_table = doc.add_table(rows=7, cols=3)
    param_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    param_table.style = 'Table Grid'

    param_headers = ["항목", "값", "단위"]
    for j, h in enumerate(param_headers):
        cell = param_table.cell(0, j)
        cell.text = h
        set_cell_bg(cell, "1E40AF")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = '맑은 고딕'

    rows_data = [
        ("궤도 장반경 (a)", f"{form_data.get('a_km', 0):,.0f}", "km"),
        ("이심률 (e)",       f"{form_data.get('e', 0):.2f}",    "—"),
        ("공전 주기 (T)",    f"{orbit['T_h']:.3f}",              "시간"),
        ("공전 주기 (T)",    f"{orbit['T_d']:.5f}",              "일"),
        ("근일점 거리 (rₚ)", f"{orbit['rp_km']:,.1f}",           "km"),
        ("원일점 거리 (rₐ)", f"{orbit['ra_km']:,.1f}",           "km"),
    ]
    for i, (label, val, unit) in enumerate(rows_data, start=1):
        param_table.cell(i, 0).text = label
        param_table.cell(i, 1).text = val
        param_table.cell(i, 2).text = unit
        if i % 2 == 0:
            for j in range(3):
                set_cell_bg(param_table.cell(i, j), "EFF6FF")

    doc.add_paragraph()

    # 속도 비교 테이블
    add_heading(doc, "2. 궤도 속도 분석", level=1, color=(30, 58, 138))
    doc.add_paragraph()

    speed_table = doc.add_table(rows=3, cols=3)
    speed_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    speed_table.style = 'Table Grid'

    speed_headers = ["구분", "속도 (km/s)", "비고"]
    for j, h in enumerate(speed_headers):
        cell = speed_table.cell(0, j)
        cell.text = h
        set_cell_bg(cell, "065F46")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = '맑은 고딕'

    speed_rows = [
        ("근일점 속도 (최대)", f"{orbit['vp_kms']:.4f}", "가장 빠른 지점"),
        ("원일점 속도 (최소)", f"{orbit['va_kms']:.4f}", "가장 느린 지점"),
    ]
    for i, (label, val, note) in enumerate(speed_rows, start=1):
        speed_table.cell(i, 0).text = label
        speed_table.cell(i, 1).text = val
        speed_table.cell(i, 2).text = note

    ratio_para = doc.add_paragraph()
    ratio_para.add_run(
        f"\n  ▶ 속도 배율: 근일점 / 원일점 = {orbit['vp_kms']/orbit['va_kms']:.3f} 배"
    ).font.bold = True

    doc.add_paragraph()

    # ── 3~6: 보고서 본문 섹션 ──────────────────────────────────
    section_titles = [
        ("3. 궤도 설계 배경 및 목적", "orbit_purpose"),
        ("4. 물리 법칙 적용 — 케플러 제3법칙 계산 과정", "kepler_calc"),
        ("5. 속도 변화의 물리적 인과관계 (케플러 제2법칙 또는 에너지 보존)", "speed_analysis"),
        ("6. 미션 적합성 평가 (관측 vs 통신 트레이드오프)", "mission_eval"),
        ("7. 결론 및 한계", "conclusion"),
    ]

    for section_num, (title, key) in enumerate(section_titles, start=1):
        add_section_divider(doc)
        doc.add_paragraph()
        add_heading(doc, title, level=1, color=(30, 58, 138))
        content = form_data.get(key, "").strip()
        if content:
            for line in content.split("\n"):
                p = doc.add_paragraph(line or " ")
                p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph("(내용 없음)")
            p.runs[0].font.color.rgb = RGBColor(180, 180, 180)
        doc.add_paragraph()

    # ── 푸터 ────────────────────────────────────────────────────
    add_section_divider(doc)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"사곡고등학교 물리학II | 수행평가 보고서 | 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(148, 163, 184)

    # ── 메모리에 저장 ────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Streamlit 페이지 본체 ──────────────────────────────────────────
def run_sim():
    st.set_page_config(page_title="연구 보고서 작성", layout="wide")

    # ── 헤더 ───────────────────────────────────────────────────
    st.markdown("""
    <div style="background:linear-gradient(135deg,#1e3a8a,#1e40af,#0f172a);
                border-radius:24px;padding:40px 48px;color:white;margin-bottom:24px;
                border:1px solid #1e40af44">
        <div style="font-size:0.7rem;font-weight:900;letter-spacing:0.25em;
                    color:#60a5fa;text-transform:uppercase;margin-bottom:12px">
            📑 Performance Assessment · KASA Research Report
        </div>
        <h1 style="font-size:2rem;font-weight:900;margin:0 0 12px 0;line-height:1.2">
            화성 탐사 궤도 설계<br>연구보고서 작성
        </h1>
        <p style="color:#bfdbfe;font-size:0.95rem;margin:0;line-height:1.7">
            아래 양식을 작성한 후 <strong>「보고서 생성」</strong> 버튼을 누르면 
            <code style="background:#1e40af;padding:2px 8px;border-radius:4px">.docx</code> 
            파일로 다운로드할 수 있습니다.
        </p>
    </div>
    """, unsafe_allow_html=True)

    if not DOCX_AVAILABLE:
        st.error("""
        ⚠️ `python-docx` 라이브러리가 설치되어 있지 않습니다.
        터미널에서 다음 명령어를 실행하세요:
        ```
        pip install python-docx
        ```
        """)

    # ── 사이드바: 궤도 파라미터 (시뮬레이터와 동기화 안내) ────────
    with st.sidebar:
        st.markdown("### 🛸 궤도 파라미터 입력")
        st.caption("🔗 시뮬레이터에서 설계한 값을 여기에 동일하게 입력하세요.")
        st.markdown("---")

        planet_name = st.selectbox("🌍 중심 천체", list(PLANETS.keys()), index=0)
        planet_info = PLANETS[planet_name]
        R_km = planet_info["R_km"]

        a_km = st.number_input(
            "궤도 장반경 a (km)",
            min_value=int(R_km) + 100,
            max_value=int(R_km) + 50000,
            value=max(int(R_km) + 3500, 7000),
            step=100,
        )
        e = st.slider("이심률 e", 0.0, 0.80, 0.30, 0.01)
        orbit = calc(a_km, e, planet_info["M"])

        st.markdown("---")
        st.markdown("#### 📊 자동 계산 결과")

        def sc(label, val, unit="", color="#38bdf8"):
            st.markdown(f"""
            <div style="background:#0f172a;border-left:3px solid {color};
                        border-radius:8px;padding:8px 12px;margin:5px 0">
                <div style="color:#64748b;font-size:0.68rem;font-weight:700">{label}</div>
                <div style="color:{color};font-size:0.95rem;font-weight:800">
                    {val} <span style="color:#475569;font-size:0.7rem">{unit}</span>
                </div>
            </div>
            """, unsafe_allow_html=True)

        sc("공전 주기 T", f"{orbit['T_h']:.3f}", "h", "#a78bfa")
        sc("공전 주기 T", f"{orbit['T_d']:.5f}", "일", "#a78bfa")
        sc("근일점", f"{orbit['rp_km']:,.0f}", "km", "#f87171")
        sc("원일점", f"{orbit['ra_km']:,.0f}", "km", "#60a5fa")
        sc("근일점 속도", f"{orbit['vp_kms']:.4f}", "km/s", "#4ade80")
        sc("원일점 속도", f"{orbit['va_kms']:.4f}", "km/s", "#fbbf24")

    # ── 메인 입력 폼 ───────────────────────────────────────────
    st.markdown("### 📝 보고서 기본 정보")
    col1, col2, col3 = st.columns(3)
    with col1:
        student_id   = st.text_input("학번", placeholder="예: 30201")
    with col2:
        student_name = st.text_input("이름", placeholder="예: 홍길동")
    with col3:
        submit_date  = st.date_input("제출일", value=datetime.today()).strftime("%Y년 %m월 %d일")

    st.markdown("---")

    # ── 섹션 1: 설계 목적 ─────────────────────────────────────
    st.markdown("""
    <div style="background:#1e293b;border-radius:16px;padding:16px 20px;
                margin-bottom:8px;border-left:4px solid #3b82f6">
        <div style="color:#93c5fd;font-size:0.75rem;font-weight:900;text-transform:uppercase;
                    letter-spacing:0.15em">Section 3</div>
        <div style="color:white;font-size:1.05rem;font-weight:800;margin-top:4px">
            궤도 설계 배경 및 목적
        </div>
        <div style="color:#64748b;font-size:0.8rem;margin-top:4px">
            왜 이 장반경과 이심률을 선택했는지, 어떤 미션(촬영/통신)을 목표로 했는지 서술하세요.
        </div>
    </div>
    """, unsafe_allow_html=True)
    orbit_purpose = st.text_area(
        "orbit_purpose", label_visibility="collapsed",
        placeholder=(
            "예시: 본 연구는 화성 표면의 고해상도 촬영을 1차 목표로 하여 궤도를 설계하였다. "
            "이심률 e = 0.35로 설정함으로써 근일점 고도를 약 xxx km로 낮추어 촬영 해상도를 "
            "최대화하였고, 동시에 원일점 고도를 xxx km 이상으로 유지하여..."
        ),
        height=150,
        key="orbit_purpose_input"
    )

    st.markdown("---")

    # ── 섹션 2: 케플러 제3법칙 계산 과정 ─────────────────────
    st.markdown("""
    <div style="background:#1e293b;border-radius:16px;padding:16px 20px;
                margin-bottom:8px;border-left:4px solid #8b5cf6">
        <div style="color:#c4b5fd;font-size:0.75rem;font-weight:900;text-transform:uppercase;
                    letter-spacing:0.15em">Section 4</div>
        <div style="color:white;font-size:1.05rem;font-weight:800;margin-top:4px">
            물리 법칙 적용 — 케플러 제3법칙 계산 과정
        </div>
        <div style="color:#64748b;font-size:0.8rem;margin-top:4px">
            T² = (4π²/GM)a³ 공식을 사용하여 공전 주기를 직접 계산하는 과정을 단계별로 서술하세요.
        </div>
    </div>
    """, unsafe_allow_html=True)

    # 계산 힌트 표시
    with st.expander("💡 계산 힌트 보기"):
        st.latex(r"T = 2\pi\sqrt{\frac{a^3}{GM}}")
        st.markdown(f"""
        - $G = 6.674 \\times 10^{{-11}}$ N·m²/kg²
        - $M_{{{planet_name}}} = {planet_info['M']:.3e}$ kg
        - $a = {a_km:,}$ km $ = {a_km*1000:.3e}$ m
        - 계산 결과: $T = {orbit['T_h']:.3f}$ h $= {orbit['T_d']:.5f}$ 일
        """)

    kepler_calc = st.text_area(
        "kepler_calc", label_visibility="collapsed",
        placeholder=(
            "예시:\n"
            "1단계: 장반경 a를 미터로 환산\n"
            "   a = 7,000 km × 1000 = 7.000 × 10⁶ m\n\n"
            "2단계: 케플러 제3법칙 공식 적용\n"
            "   T = 2π√(a³ / GM)\n"
            "   T = 2π√((7.000×10⁶)³ / (6.674×10⁻¹¹ × 6.39×10²³))\n\n"
            "3단계: 계산 결과\n"
            "   T = ... 초 = ... 시간"
        ),
        height=200,
        key="kepler_calc_input"
    )

    st.markdown("---")

    # ── 섹션 3: 속도 변화 분석 ───────────────────────────────
    st.markdown("""
    <div style="background:#1e293b;border-radius:16px;padding:16px 20px;
                margin-bottom:8px;border-left:4px solid #10b981">
        <div style="color:#6ee7b7;font-size:0.75rem;font-weight:900;text-transform:uppercase;
                    letter-spacing:0.15em">Section 5</div>
        <div style="color:white;font-size:1.05rem;font-weight:800;margin-top:4px">
            속도 변화의 물리적 인과관계
        </div>
        <div style="color:#64748b;font-size:0.8rem;margin-top:4px">
            근일점과 원일점에서 속도가 다른 이유를 케플러 제2법칙 또는 역학적 에너지 보존으로 설명하세요.
        </div>
    </div>
    """, unsafe_allow_html=True)

    speed_analysis = st.text_area(
        "speed_analysis", label_visibility="collapsed",
        placeholder=(
            "예시: 근일점에서의 속도가 원일점보다 큰 이유는 두 가지 관점에서 설명할 수 있다.\n\n"
            "[케플러 제2법칙 관점]\n"
            "단위 시간당 쓸리는 면적이 일정하므로, 천체와 가까울수록(r이 작을수록)\n"
            "속도가 빨라야 같은 면적을 쓸 수 있다...\n\n"
            "[역학적 에너지 보존 관점]\n"
            "역학적 에너지 E = KE + PE = (1/2)mv² - GMm/r = const\n"
            "근일점에서 r이 최소이므로 PE가 가장 음수이고, KE가 최대가 된다..."
        ),
        height=200,
        key="speed_analysis_input"
    )

    st.markdown("---")

    # ── 섹션 4: 미션 적합성 평가 ─────────────────────────────
    st.markdown("""
    <div style="background:#1e293b;border-radius:16px;padding:16px 20px;
                margin-bottom:8px;border-left:4px solid #f59e0b">
        <div style="color:#fcd34d;font-size:0.75rem;font-weight:900;text-transform:uppercase;
                    letter-spacing:0.15em">Section 6</div>
        <div style="color:white;font-size:1.05rem;font-weight:800;margin-top:4px">
            미션 적합성 평가
        </div>
        <div style="color:#64748b;font-size:0.8rem;margin-top:4px">
            설계한 궤도가 '고해상도 촬영'과 '광범위 통신' 목적에 얼마나 적합한지 수치를 근거로 평가하세요.
        </div>
    </div>
    """, unsafe_allow_html=True)

    rp_alt = orbit["rp_km"] - PLANETS[planet_name]["R_km"]
    ra_alt = orbit["ra_km"] - PLANETS[planet_name]["R_km"]
    st.info(f"📊 현재 근일점 고도: **{rp_alt:.0f} km** | 원일점 고도: **{ra_alt:.0f} km** | 주기: **{orbit['T_h']:.2f} h**")

    mission_eval = st.text_area(
        "mission_eval", label_visibility="collapsed",
        placeholder=(
            "예시: 설계된 궤도의 근일점 고도는 xxx km로, 고해상도 촬영에 있어 [적합/부적합]하다.\n"
            "이유는 ...\n\n"
            "원일점 고도는 xxx km로, 화성 표면 커버리지 측면에서 [우수/미흡]하다.\n"
            "통신 지연을 고려할 때 공전 주기 xxx 시간은 [허용 가능/지나치게 길다]..."
        ),
        height=180,
        key="mission_eval_input"
    )

    st.markdown("---")

    # ── 섹션 5: 결론 ──────────────────────────────────────────
    st.markdown("""
    <div style="background:#1e293b;border-radius:16px;padding:16px 20px;
                margin-bottom:8px;border-left:4px solid #ef4444">
        <div style="color:#fca5a5;font-size:0.75rem;font-weight:900;text-transform:uppercase;
                    letter-spacing:0.15em">Section 7</div>
        <div style="color:white;font-size:1.05rem;font-weight:800;margin-top:4px">
            결론 및 한계
        </div>
        <div style="color:#64748b;font-size:0.8rem;margin-top:4px">
            연구를 통해 얻은 결론과 이 모델의 한계(원 궤도 근사, 대기 저항 무시 등)를 서술하세요.
        </div>
    </div>
    """, unsafe_allow_html=True)

    conclusion = st.text_area(
        "conclusion", label_visibility="collapsed",
        placeholder=(
            "예시: 본 연구에서는 케플러 제3법칙을 이용하여 화성 탐사선의 공전 주기를 이론적으로 계산하고,\n"
            "시뮬레이터를 통해 검증하였다. 이심률이 커질수록 근일점 속도와 원일점 속도의 차이가 증가하며,\n"
            "이는 역학적 에너지 보존과 케플러 제2법칙 모두로 설명 가능하다.\n\n"
            "[연구의 한계]\n"
            "  ① 화성의 대기 저항을 고려하지 않았다.\n"
            "  ② 화성의 자전에 의한 조석 효과를 무시하였다.\n"
            "  ③ 태양 복사압 등 섭동력은 계산에 포함되지 않았다."
        ),
        height=180,
        key="conclusion_input"
    )

    # ── 보고서 생성 버튼 ──────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📥 보고서 생성 및 다운로드")

    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("🗒️ 연구보고서 DOCX 생성", use_container_width=True, type="primary"):
            if not student_name.strip():
                st.error("⚠️ 이름을 입력해주세요.")
            elif not student_id.strip():
                st.error("⚠️ 학번을 입력해주세요.")
            elif not DOCX_AVAILABLE:
                st.error("⚠️ python-docx가 설치되지 않아 파일을 생성할 수 없습니다.")
            else:
                form_data = {
                    "student_id":    student_id,
                    "student_name":  student_name,
                    "submit_date":   submit_date,
                    "planet":        planet_name,
                    "a_km":          a_km,
                    "e":             e,
                    "orbit_purpose": orbit_purpose,
                    "kepler_calc":   kepler_calc,
                    "speed_analysis":speed_analysis,
                    "mission_eval":  mission_eval,
                    "conclusion":    conclusion,
                }
                try:
                    docx_bytes = create_docx(form_data, orbit)
                    filename = f"다누리2호_궤도보고서_{student_id}_{student_name}.docx"
                    st.download_button(
                        label="⬇️ DOCX 파일 다운로드",
                        data=docx_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    st.success(f"✅ 보고서가 생성되었습니다! 위 버튼을 눌러 다운로드하세요.")
                    st.balloons()
                except Exception as ex:
                    st.error(f"보고서 생성 중 오류가 발생했습니다: {ex}")

    st.markdown("---")
    st.caption("📌 사곡고등학교 물리학II · 수행평가 지원 시스템 v1.0")


if __name__ == "__main__":
    run_sim()
