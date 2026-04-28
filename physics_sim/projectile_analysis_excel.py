import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn # 한글 폰트 설정을 위한 qn 임포트 추가
import io

# 페이지 설정

# CSS를 활용한 엑셀 스타일 및 인쇄 최적화
st.markdown("""
    <style>
    /* 엑셀 스타일 디자인 */
    .excel-header { background-color: #3b82f6; color: white; padding: 5px; text-align: center; font-weight: bold; border: 1px solid #ddd; }
    .excel-value { background-color: #f3f4f6; padding: 5px; text-align: center; border: 1px solid #ddd; color: black !important; }
    .result-header { background-color: #fecaca; color: black; padding: 5px; text-align: center; font-weight: bold; border: 1px solid #ddd; }
    .result-value { background-color: #fee2e2; padding: 5px; text-align: center; border: 1px solid #ddd; color: black !important; }

    .print-header { font-size: 1.1rem; font-weight: bold; margin-bottom: 15px; border-bottom: 2px solid black; padding-bottom: 8px; color: inherit; }
    </style>
""", unsafe_allow_html=True)

st.title("📊 포물선 운동 정밀 데이터 분석")

# 📝 탐구 질문 및 데이터 분석
st.markdown("""
1. 초기 속도가 같을 때 가장 멀리 날아갈 수 있는 각도는?
2. 초기 속도가 같을 때 수평 도달 거리가 같은 각도들은 어떤 관계인가?
3. 초기 속도가 같을 때 수평도달 거리가 같은 발사각에서 최고점 도달 시간, 최고점의 높이는 어떤 차이가 있는가?
4. 수평 도달 거리 $R$과 최고점 높이 $H$ 사이에는 어떤 정량적 관계가 성립하는가?
""")

# --- 상단 입력부 및 주요 결과값 ---
col_in1, col_in2, col_res = st.columns([1, 1, 2])

with col_in1:
    theta_deg = st.number_input("발사각 (θ, 도)", min_value=0.0, max_value=90.0, value=30.0, step=1.0)
    v0 = st.number_input("초기 속도 (v₀, m/s)", min_value=1.0, max_value=100.0, value=40.0, step=1.0)

theta = np.radians(theta_deg)
g = st.radio("🌍 중력 가속도 g (m/s²)", options=[9.8, 10.0], index=0, horizontal=True)

# 물리량 계산
vx0 = v0 * np.cos(theta)
vy0 = v0 * np.sin(theta)
t_H = vy0 / g  # 최고점 도달 시간
H = (vy0**2) / (2 * g)  # 최고점 높이
t_R = 2 * t_H  # 지면 도달 시간
R = vx0 * t_R  # 수평 도달 거리

with col_in2:
    st.markdown(f'<div class="excel-header">x방향 처음 속도</div><div class="excel-value">{vx0:.2f}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="excel-header">y방향 처음 속도</div><div class="excel-value">{vy0:.2f}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="excel-header">중력 가속도</div><div class="excel-value">{g}</div>', unsafe_allow_html=True)

with col_res:
    r_c1, r_c2 = st.columns(2)
    with r_c1:
        st.markdown('<div class="result-header">최고점 도달시간</div>', unsafe_allow_html=True)
        st.markdown('<div class="result-header">최고점의 높이</div>', unsafe_allow_html=True)
        st.markdown('<div class="result-header">수평도달 거리</div>', unsafe_allow_html=True)
    with r_c2:
        st.markdown(f'<div class="result-value">{t_H:.2f}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="result-value">{H:.2f}</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="result-value">{R:.2f}</div>', unsafe_allow_html=True)

st.divider()

# --- 데이터 생성 (0.1초 단위) ---
t_steps = np.arange(0, t_R + 0.1, 0.1)
x_vals = vx0 * t_steps
y_vals = np.maximum(vy0 * t_steps - 0.5 * g * t_steps**2, 0)
vx_vals = np.full_like(t_steps, vx0)
vy_vals = vy0 - g * t_steps
ay_vals = np.full_like(t_steps, -g)

df = pd.DataFrame({
    "시간(t)": t_steps,
    "x위치(m)": x_vals,
    "y위치(m)": y_vals,
    "v_x(속도)": vx_vals,
    "v_y(속도)": vy_vals,
    "a_y(가속도)": ay_vals
})

# --- [데이터 리포트 내보내기] ---
st.divider()
st.subheader("📤 결과 데이터 리포트 저장")

# --- Word 보고서 생성 함수 ---
def create_docx():
    doc = Document()
    # 기본 폰트 설정 (한글 깨짐 방지)
    doc.styles['Normal'].font.name = '맑은 고딕'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    title = doc.add_heading('포물선 운동 실험 결과 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    # 기본 정보
    doc.add_paragraph(' 학년: ____   반: ____   번호: ____   이름: __________')
    doc.add_paragraph(f'실험 조건: 발사각 {theta_deg}°, 초기속도 {v0}m/s, 중력가속도 {g}m/s²')
    
    # 탐구 질문 및 답변
    doc.add_heading('📝 탐구 질문 및 답변', level=1)
    questions = [
        "1. 초기 속도가 같을 때 가장 멀리 날아갈 수 있는 각도는?",
        "2. 초기 속도가 같을 때 수평 도달 거리가 같은 각도들은 어떤 관계인가?",
        "3. 초기 속도가 같을 때 수평도달 거리가 같은 발사각에서 최고점 도달 시간, 최고점의 높이는 어떤 차이가 있는가?",
        "4. 수평 도달 거리 R과 최고점 높이 H 사이에는 어떤 정량적 관계가 성립하는가?"
    ]
    for q in questions:
        p_q = doc.add_paragraph(q)
        for run in p_q.runs:
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        doc.add_paragraph("\n" * 3)
        doc.add_paragraph("_" * 50)
        
    # 결과 요약
    head_res = doc.add_heading('📊 실험 결과 요약', level=1)
    for run in head_res.runs:
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '구분'
    hdr_cells[1].text = '결괏값'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '맑은 고딕'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    
    results = [
        ("최고점 도달 시간", f"{t_H:.2f} s"), ("최고점 높이", f"{H:.2f} m"), ("수평 도달 거리", f"{R:.2f} m")
    ]
    for item, val in results:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = val
        
    doc.add_paragraph("\n* 위 데이터는 시뮬레이션 기반 결과입니다.")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

col_ex1, col_ex2 = st.columns(2)
with col_ex1:
    docx_data = create_docx()
    st.download_button("📝 보고서 양식 다운로드 (DOCX)", data=docx_data, file_name=f"projectile_report_{theta_deg}deg.docx", mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', use_container_width=True)
with col_ex2:
    csv = df.to_csv(index=False).encode('utf-8-sig')
    st.download_button("📂 실험 데이터 다운로드 (CSV/Excel용)", data=csv, file_name=f"projectile_data_{theta_deg}deg.csv", mime='text/csv', use_container_width=True)

# --- 메인 레이아웃: 표(좌) + 그래프(우) ---
col_table, col_graphs = st.columns([2, 3])

with col_table:
    st.subheader("📋 데이터 시트")
    st.dataframe(df.style.format("{:.2f}"), height=600, use_container_width=True)

with col_graphs:
    st.subheader("📈 물리량 그래프 분석")
    
    # 2x2 서브플롯 생성
    fig = make_subplots(
        rows=2, cols=2,
        subplot_titles=("시간 - x위치", "시간 - y위치", "시간 - vx", "시간 - vy"),
        vertical_spacing=0.15,
        horizontal_spacing=0.1
    )

    # Graph 1: x-t
    fig.add_trace(go.Scatter(x=t_steps, y=x_vals, mode='lines+markers', name='x(t)', line=dict(color='blue')), row=1, col=1)
    # Graph 2: y-t
    fig.add_trace(go.Scatter(x=t_steps, y=y_vals, mode='lines+markers', name='y(t)', line=dict(color='red')), row=1, col=2)
    # Graph 3: vx-t
    fig.add_trace(go.Scatter(x=t_steps, y=vx_vals, mode='lines+markers', name='vx(t)', line=dict(color='green')), row=2, col=1)
    # Graph 4: vy-t
    fig.add_trace(go.Scatter(x=t_steps, y=vy_vals, mode='lines+markers', name='vy(t)', line=dict(color='orange')), row=2, col=2)

    # 그래프 스타일 설정 (테마에 맞는 템플릿 사용)
    p_template = "plotly_dark" if st.get_option("theme.base") == "dark" else "plotly_white"
    
    fig.update_layout(
        height=600, 
        showlegend=False, 
        margin=dict(l=20, r=20, t=40, b=20),
        template=p_template,
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)'
    )
    st.plotly_chart(fig, use_container_width=True)

st.info("💡 각 수치를 변경하면 표와 그래프가 실시간으로 업데이트됩니다.")
