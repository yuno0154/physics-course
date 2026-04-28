import streamlit as st
import pandas as pd
import numpy as np
import json
import base64
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn 

# --- 로드 대기 데이터 처리 (위젯 렌더링 전에 반드시 실행) ---
if '_pending_load' in st.session_state:
    pending = st.session_state.pop('_pending_load')
    for k, v in pending.items():
        st.session_state[k] = v

# ── Session State 초기화 ────────────────────────────────────────────
if "report_content" not in st.session_state:
    st.session_state["report_content"] = ""

# CSS를 활용한 인쇄 최적화 및 폰트 크기 설정
st.markdown("""
    <style>
    /* 메인 콘텐츠 영역에만 폰트 설정 적용 (사이드바 제외) */
    [data-testid="stAppViewMainContent"] {
        font-size: 12pt !important;
    }

    /* 제목 크기 (18pt) */
    [data-testid="stAppViewMainContent"] h1 {
        font-size: 18pt !important;
        margin-bottom: 0.5rem !important;
    }

    /* 주요 섹션 제목 (13pt) */
    [data-testid="stAppViewMainContent"] h2,
    [data-testid="stAppViewMainContent"] h3 {
        font-size: 13pt !important;
        margin-top: 1rem !important;
        margin-bottom: 0.5rem !important;
    }

    /* 질문 및 소제목 (12pt Bold) */
    [data-testid="stAppViewMainContent"] h4,
    [data-testid="stAppViewMainContent"] h5,
    [data-testid="stAppViewMainContent"] h6 {
        font-size: 12pt !important;
        font-weight: bold !important;
    }

    @media print {
        @page { margin: 15mm; size: A4; }
        header, [data-testid="stSidebar"], [data-testid="stToolbar"], .stActionButton, [data-testid="stExpander"], [data-testid="stFileUploader"], button { display: none !important; }
        .main .block-container { padding: 0 !important; }
        
        /* 겹침 방지 */
        .stMarkdown, .stTable, .stPlotlyChart, .stImage, section, div.row-widget { 
            page-break-inside: avoid !important; 
            display: block !important;
            margin-bottom: 1.5rem !important;
            float: none !important;
            position: relative !important;
        }
        img { max-width: 100% !important; height: auto !important; }
    }
    .answer-box { border: 1px solid #ccc; min-height: 80px; padding: 10px; margin-bottom: 20px; background-color: #f9f9f9; }
    </style>
""", unsafe_allow_html=True)

# --- 이미지 파일을 base64로 인코딩하는 함수 ---
def file_to_b64(file_obj):
    """업로드된 파일 객체를 base64 문자열로 변환"""
    if file_obj is None:
        return None, None
    try:
        file_obj.seek(0)
        b64 = base64.b64encode(file_obj.read()).decode('utf-8')
        return b64, file_obj.name
    except Exception:
        return None, None

def get_img_b64_for_save(key):
    """저장 시 이미지 b64 데이터 확보: 새 업로드 파일 우선, 없으면 기존 session_state 사용"""
    file_obj = st.session_state.get(key)
    if file_obj is not None:
        b64, name = file_to_b64(file_obj)
        if b64:
            return b64, name
    # 새 파일 없으면 기존 저장된 b64 재사용
    return st.session_state.get(f'{key}_b64'), st.session_state.get(f'{key}_name')

st.divider()

st.divider()

# --- [상단 고정] 학생 정보 및 데이터 관리 ---
st.markdown("### 👤 학생 정보")
info_c1, info_c2, info_c3 = st.columns(3)
with info_c1:
    class_num = st.text_input("반 (Class)", placeholder="예: 3-1", key="class_num")
with info_c2:
    student_num = st.text_input("번호 (Number)", placeholder="예: 01", key="student_num")
with info_c3:
    student_name = st.text_input("성함 (Name)", placeholder="홍길동", key="student_name")

# 답변/데이터 변수 사전 로드 (NameError 방지 및 보고서 생성용)
qa_data = {f"a{i}": st.session_state.get(f"a{i}", "") for i in range(1, 6)}
m_inputs = {f"m_{i}": st.session_state.get(f"m_input_{i}", "") for i in range(3)}
b_theories = {f"b_{i}": st.session_state.get(f"b_theory_{i}", "") for i in range(3)}

st.divider()

# --- 상단 헤더 섹션 (제목 및 작업 버튼 한 줄 배치) ---
header_col1, header_col2, header_col3, header_col4 = st.columns([3.5, 0.8, 0.8, 0.8])

with header_col1:
    st.markdown("<h2 style='margin:0; padding:0; line-height:1.5;'>📹 [수행평가 1-1] 영상 분석 보고서</h2>", unsafe_allow_html=True)

with header_col2:
    if st.button("🖨️ 인쇄", key="print_btn", use_container_width=True):
        st.components.v1.html("<script>parent.window.print()</script>", height=0)

with header_col3:
    with st.popover("📥 로드", use_container_width=True):
        uploaded_json = st.file_uploader("JSON 파일 선택", type=["json"], key="load_json", label_visibility="collapsed")
        if uploaded_json:
            try:
                loaded_data = json.load(uploaded_json)
                st.session_state['_pending_load'] = loaded_data
                st.rerun()
            except Exception as e:
                st.error(f"오류: {e}")

with header_col4:
    # --- Word 보고서 생성 함수 ---
    def create_eval_docx():
        doc = Document()
        doc.styles['Normal'].font.name = '맑은 고딕'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        
        title = doc.add_heading('[수행평가 1-1] 영상 분석 및 데이터 해석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        doc.add_paragraph(f"반: {class_num}  번호: {student_num}  이름: {student_name}")
        
        # 1. 시각적 자료 (이미지)
        doc.add_heading('📊 실험 데이터 및 그래프', level=1)
        for img_key, label in [('data_file', '📋 분석 데이터 표'), ('g_img_pos', '① 시간-위치 그래프'), ('g_img_vel', '② 시간-속도 그래프')]:
            img_file = st.session_state.get(img_key)
            img_b64 = st.session_state.get(f"{img_key}_b64")
            
            if img_file or img_b64:
                doc.add_heading(label, level=2)
                try:
                    img_data = img_file if img_file else io.BytesIO(base64.b64decode(img_b64))
                    if img_key == 'data_file' and hasattr(img_data, 'name') and img_data.name.lower().endswith('.csv'):
                        doc.add_paragraph("(CSV 데이터 파일은 이미지로 포함되지 않습니다.)")
                    else:
                        doc.add_picture(img_data, width=Inches(5.5))
                except:
                    doc.add_paragraph(f"({label} 이미지 로드 오류)")

        doc.add_page_break()

        # 2. 분석 및 답변
        doc.add_heading('🤔 탐구 결과 정리 및 분석', level=1)
        qa_labels = [
            "가. 수평 방향 운동에서 속력은 시간에 따라 어떻게 변하는가?",
            "나. 수평 방향 운동이 그렇게 일어나는 이유는?",
            "다. 연직 방향의 운동에서 속력은 시간에 따라 어떻게 변하는가?",
            "라. 연직 방향 운동이 그렇게 일어나는 이유는?"
        ]
        for i, label in enumerate(qa_labels, 1):
            p_q = doc.add_paragraph(f"{label}", style='List Bullet')
            for run in p_q.runs: run.font.bold = True
            doc.add_paragraph(st.session_state.get(f"a{i}", "(답변 없음)"))

        # 3. 측정 데이터 표
        doc.add_heading('📏 측정 결과 및 비교', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        cells = table.rows[0].cells
        for i, h in enumerate(["항목", "실측값", "이론값"]): cells[i].text = h
        
        m_labels = ["최고점 도달 시간(s)", "최고점 높이(m)", "수평도달 거리(m)"]
        for i in range(3):
            row = table.add_row().cells
            row[0].text = m_labels[i]
            row[1].text = st.session_state.get(f"m_input_{i}", "")
            row[2].text = st.session_state.get(f"b_theory_{i}", "")

        doc.add_heading('마. 결과 비교 및 추론', level=2)
        doc.add_paragraph(st.session_state.get("a5", "(답변 없음)"))

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    docx_data = create_eval_docx()
    st.download_button(
        "📝 DOCX", 
        data=docx_data, 
        file_name=f"수행평가_보고서_{student_name}.docx", 
        key="docx_eval_btn",
        use_container_width=True
    )

# --- 탐구 기본 정보 ---
st.markdown("### 📝 탐구 개요")
c1, c2 = st.columns(2)
with c1:
    st.info("**탐구 주제**: 운동 분석 프로그램으로 운동 분석")
with c2:
    st.success("**탐구 목표**: 비스듬히 위로 던진 물체의 운동을 분석하고 뉴턴의 운동 법칙으로 설명할 수 있다.")

st.markdown("---")

# --- 준비물 및 과정 ---
st.subheader("🛠️ 준비물 및 탐구 과정")
st.markdown("""
**[준비물]**
- **동영상 분석 프로그램**: [웹 서비스 접속 (https://videoanalysis.app?key=tvavTYNBRc)](https://videoanalysis.app?key=tvavTYNBRc)
- **분석용 영상**: 야구공 또는 농구공의 포물선 운동 영상

**[탐구 과정]**
1. **프로그램 접속**: 상단 링크를 통해 동영상 분석 프로그램에 접속합니다.
2. **영상 재생 및 설정**:
    - 분석할 영상을 프로그램에서 재생시킵니다.
    - 운동 분석을 위해 **클립 설정, 교정막대자 설정, 좌표축 설정, 질점 설정**을 수행합니다.
3. **데이터 추출**: 물체의 위치 변화를 추적하여 데이터와 그래프를 생성합니다.
""")

st.divider()

# --- 탐구 결과 1: 데이터 및 그래프 ---
st.subheader("📊 탐구 결과 1: 데이터 및 그래프 기록")
st.write("동영상 분석 프로그램에서 추출한 데이터 표와 5가지 종류의 그래프를 각각 캡처하여 업로드하세요.")

# 1. 데이터 표 업로드
st.markdown("#### 📋 1. 분석 데이터 표 기록")
data_file = st.file_uploader("데이터 표 이미지 또는 CSV 파일 업로드", type=['png', 'jpg', 'jpeg', 'csv'], key="data_file")

if data_file:
    if data_file.name.lower().endswith('.csv'):
        try:
            df = pd.read_csv(data_file)
            st.dataframe(df, use_container_width=True)
            st.caption("[기록] 분석 데이터 (CSV)")
        except Exception as e:
            st.error(f"CSV 파일을 읽는데 실패했습니다: {e}")
    else:
        st.image(data_file, use_container_width=True, caption="[기록] 분석 데이터 표 이미지")
elif st.session_state.get('data_file_b64'):
    # 저장된 b64 이미지 복원
    b64 = st.session_state['data_file_b64']
    name = st.session_state.get('data_file_name', '')
    img_bytes = base64.b64decode(b64)
    if name.lower().endswith('.csv'):
        try:
            df = pd.read_csv(io.BytesIO(img_bytes))
            st.dataframe(df, use_container_width=True)
            st.caption("[기록] 분석 데이터 (CSV) — 저장 파일에서 복원됨")
        except Exception as e:
            st.error(f"CSV 복원 실패: {e}")
    else:
        st.image(img_bytes, use_container_width=True, caption="[기록] 분석 데이터 표 이미지 — 저장 파일에서 복원됨")
else:
    st.info("실습 프로그램에서 다운로드한 CSV 파일 또는 캡처한 데이터 표 이미지를 업로드해 주세요.")

st.divider()

# 2. 분석 그래프 기록
st.markdown("#### 📈 2. 분석 그래프 기록")
st.caption("실습 프로그램의 그래프 화면을 캡처하여 업로드해 주세요.")

graph_sections = [
    ("pos", "① 시간-위치 (x-t, y-t) 그래프"),
    ("vel", "② 시간-속도 (vx-t, vy-t) 그래프")
]

for key, label in graph_sections:
    with st.container():
        st.write(f"**{label}**")
        g_img = st.file_uploader(f"{label} 이미지 업로드", type=['png', 'jpg', 'jpeg'], key=f"g_img_{key}", label_visibility="collapsed")
        if g_img:
            st.image(g_img, use_container_width=True, caption=label)
        elif st.session_state.get(f'g_img_{key}_b64'):
            b64 = st.session_state[f'g_img_{key}_b64']
            img_bytes = base64.b64decode(b64)
            st.image(img_bytes, use_container_width=True, caption=f"{label} — 저장 파일에서 복원됨")
        else:
            st.info(f"{label} 이미지를 업로드해 주세요.")
        st.write("")  # 간격 조절

st.divider()

# --- 탐구 결과 2: 분석 결과 질문 ---
st.subheader("🤔 탐구 결과 2: 분석 결과 정리")
st.write("분석 데이터를 바탕으로 다음 질문에 답해 보세요.")

# 질문 및 답변 입력 영역
st.markdown("#### 가. 수평 방향 운동에서 속력은 시간에 따라 어떻게 변하는가?")
a1 = st.text_area("답변 입력 (가)", placeholder="실험 데이터를 통해 관찰한 내용을 적어주세요.", height=100, label_visibility="collapsed", key="eval_a1")

st.markdown("#### 나. 수평 방향 운동이 가와 같이 일어나는 이유는 무엇인가?")
a2 = st.text_area("답변 입력 (나)", placeholder="뉴턴의 운동 법칙을 적용하여 설명하세요.", height=100, label_visibility="collapsed", key="eval_a2")

st.markdown("#### 다. 연직 방향의 운동에서 속력은 시간에 따라 어떻게 변하는가?")
a3 = st.text_area("답변 입력 (다)", placeholder="최고점 도달 전후의 속력 변화를 포함하여 설명하세요.", height=100, label_visibility="collapsed", key="eval_a3")

st.markdown("#### 라. 연직 방향 운동에서 다와 같이 일어나는 이유는 무엇인가?")
a4 = st.text_area("답변 입력 (라)", placeholder="작용하는 힘(알짜힘)의 관점에서 설명하세요.", height=100, label_visibility="collapsed", key="eval_a4")

# --- 추가 질문 마, 바 ---
st.markdown("#### 마. 분석 데이터(표, 그래프)에서 최고점 도달 시간, 최고점의 높이, 수평 도달 거리를 찾아 기록하시오.")
m_cols = st.columns(3)
m_labels = ["최고점의 도달 시간(s)", "최고점의 높이(h)", "수평도달 거리(m)"]
for i, label in enumerate(m_labels):
    with m_cols[i]:
        st.markdown(f"<div style='background-color:#fce4d6; padding:5px; text-align:center; font-weight:bold; border:1px solid #000; font-size:12pt; color:black;'>{label}</div>", unsafe_allow_html=True)
        st.text_input(label, label_visibility="collapsed", key=f"m_input_{i}")

st.write("")  # 간격

st.markdown("#### 바. [포물선 운동 정밀 데이터 분석]에서 최고점 도달 시간을 기준으로 이론적으로 계산한 수평도달 거리와 최고점의 높이를 야구공의 값과 비교하고, 차이가 나는 이유를 추론해보자.")
st.markdown("""
<div style="margin-bottom: 20px;">
    <a href="?page=physics_sim/projectile_analysis_excel" target="_blank" style="text-decoration: none;">
        <button style="width: 100%; padding: 10px; background-color: #4CAF50; color: white; border: none; border-radius: 5px; cursor: pointer; font-weight: bold;">
            🚀 [포물선 운동 정밀 데이터 분석] 페이지 새 창에서 열기
        </button>
    </a>
    <p style="font-size: 10pt; color: #666; margin-top: 5px;">※ 현재 페이지에서 이동하면 입력 중인 데이터가 사라질 수 있으므로, 반드시 새 창에서 확인하세요.</p>
</div>
""", unsafe_allow_html=True)
st.write("**[이론값]**")
b_cols = st.columns(3)
for i, label in enumerate(m_labels):
    with b_cols[i]:
        st.markdown(f"<div style='background-color:#fce4d6; padding:5px; text-align:center; font-weight:bold; border:1px solid #000; font-size:12pt; color:black;'>{label}</div>", unsafe_allow_html=True)
        st.text_input(label + " (이론)", label_visibility="collapsed", key=f"b_theory_{i}")

st.write("**[추론 및 결과 비교 기록]**")
a5 = st.text_area("결과 비교 및 추론", placeholder="이론값과 실제값의 차이가 발생하는 원인(공기 저항 등)을 물리적 원리와 함께 추론하여 적어주세요.", height=150, label_visibility="collapsed", key="a5")

# --- 과제 제출 안내 (인쇄 제외) ---
with st.expander("📤 과제 제출 방법 안내 (인쇄 시 출력되지 않음)", expanded=True):
    st.markdown("""
    1. 작성한 결과물을 상단 **'🖨️ 보고서 인쇄 / PDF 저장'** 버튼을 눌러 PDF로 저장하세요.
    2. 파일명을 **'이름.pdf'** 또는 **'이름_이름.pdf'**로 저장하세요. (예: 홍길동_김영철.pdf)
    3. 다음 두 파일을 **TEAMS 과제** 탭에 업로드하여 제출하세요.
       - **보고서 제출** (PDF)
       - **영상 분석 파일 제출** (예: 홍길동_김영철.vmbl)
    """)

st.divider()
st.caption("사곡고등학교 물리학 II 시뮬레이션 및 실습 지원 포털 | 본 보고서는 작성 후 상단 인쇄 버튼을 통해 PDF로 저장할 수 있습니다.")
